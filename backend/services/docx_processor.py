"""
services/docx_processor.py
Word document (.docx) extraction using python-docx.
Extracts paragraphs, headings, tables. Stores in vector memory.
"""
import os, re


def process_docx(file_path: str, file_id: str, filename: str,
                 session_id: str = None, project_id: str = None) -> dict:
    if not os.path.exists(file_path):
        return {"success": False, "error": "File not found"}
    try:
        from docx import Document
    except ImportError:
        return {
            "success": False,
            "error":   "python-docx not installed. Run: pip install python-docx --break-system-packages"
        }
    try:
        doc   = Document(file_path)
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Extract tables
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                rows.append(",".join(cells))
            if rows:
                tables.append("\n".join(rows))

        full_text = "\n\n".join(paras)
        if tables:
            full_text += "\n\n" + "\n\n".join(tables)

        # Store in vector memory
        try:
            from services.vector_memory import store_document
            ns = project_id or session_id or file_id
            store_document(file_id, full_text, ns, filename, "docx")
        except Exception: pass

        return {
            "success":   True,
            "file_id":   file_id,
            "filename":  filename,
            "paragraphs":len(paras),
            "tables":    len(tables),
            "chars":     len(full_text),
            "preview":   full_text[:400],
            "table_data":tables[:2],
        }
    except Exception as e:
        return {"success": False, "error": str(e)}
